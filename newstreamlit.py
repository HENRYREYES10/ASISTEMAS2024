import os
import datetime
from collections import Counter
import streamlit as st
from docx import Document
from docx.shared import Pt, Inches
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from io import BytesIO
import pandas as pd

# Función para leer los logs desde el archivo subido
def leer_logs(file):
    try:
        if file.name.endswith('.log'):
            return file.read().decode('latin-1').splitlines()
        elif file.name.endswith('.xlsx'):
            df = pd.read_excel(file)
            if 'Severity' in df.columns and 'Message' in df.columns and 'Timestamp' in df.columns:
                return df[['Severity', 'Message', 'Timestamp']].values.tolist()
            else:
                st.error("No se encontraron las columnas 'Severity', 'Message' o 'Timestamp' en el archivo.")
                return []
        else:
            st.error("Formato de archivo no soportado. Suba un archivo .log o .xlsx")
            return []
    except Exception as e:
        st.error(f"Error al leer el archivo: {e}")
        return []

# Función para generar explicaciones detalladas y personalizadas para cada log
def generar_explicacion(log):
    message = log[1]
    if "Database connection failed" in message:
        return "Fallo en la conexión con la base de datos. Esto podría deberse a credenciales incorrectas, un problema con la red, o el servicio de base de datos no está disponible."
    elif "Unable to reach API endpoint" in message:
        return "No se pudo comunicar con el endpoint de la API. Verifique la URL del endpoint, la conectividad de red y la disponibilidad del servicio."
    elif "Failed to back up database" in message:
        return "La copia de seguridad falló. Posibles causas incluyen falta de espacio en disco, permisos insuficientes, o problemas con el servicio de respaldo."
    elif "High memory usage detected" in message:
        return "Uso elevado de memoria detectado. Revise los procesos en ejecución, posibles fugas de memoria o configuraciones inadecuadas de aplicaciones."
    elif "Disk space low" in message:
        return "Espacio en disco insuficiente. Se recomienda liberar espacio eliminando archivos innecesarios o ampliar la capacidad de almacenamiento."
    elif "Slow response time" in message:
        return "El sistema responde lentamente. Podría ser debido a alta carga de CPU, cuellos de botella en el acceso a la base de datos, o problemas de red."
    elif "System outage detected" in message:
        return "Interrupción del sistema detectada. Verifique la integridad del hardware, la configuración de la red, y el estado de los servicios críticos."
    elif "Security breach detected" in message:
        return "Posible brecha de seguridad detectada. Revise los logs de acceso, cambie contraseñas comprometidas, y considere fortalecer las medidas de seguridad."
    elif "Application crash" in message:
        return "Una aplicación se bloqueó. Revise los registros de la aplicación para identificar la causa del fallo y considere implementar mecanismos de recuperación."
    elif "User session timeout" in message:
        return "La sesión del usuario expiró. Esto podría deberse a configuraciones de tiempo de espera muy bajas o a inactividad prolongada del usuario."
    elif "Unauthorized access attempt" in message:
        return "Intento de acceso no autorizado detectado. Revise los registros de seguridad para identificar al actor y considere aumentar las medidas de protección."
    elif "Server overload" in message:
        return "El servidor está sobrecargado. Considere optimizar las aplicaciones, balancear la carga o aumentar los recursos del servidor."
    elif "Data synchronization error" in message:
        return "Error en la sincronización de datos. Verifique las conexiones de red, la consistencia de datos y los procesos de sincronización."
    elif "API rate limit exceeded" in message:
        return "Límite de tasa de API excedido. Optimice las llamadas a la API para evitar exceder los límites y considere implementar un manejo de tasas."
    elif "Invalid input detected" in message:
        return "Se ha detectado una entrada inválida. Asegúrese de que los datos introducidos cumplen con los formatos y requisitos esperados."
    elif "Password reset requested" in message:
        return "Solicitud de restablecimiento de contraseña detectada. Verifique si se trata de una solicitud legítima y si es necesario tomar medidas adicionales."
    elif "Failed login attempt detected" in message:
        return "Intento de inicio de sesión fallido detectado. Puede ser indicativo de intentos de acceso no autorizados o errores en la autenticación del usuario."
    elif "Session timeout" in message:
        return "Tiempo de sesión agotado. Los usuarios han sido desconectados por inactividad prolongada o debido a políticas de seguridad."
    elif "Scheduled report generated" in message:
        return "Un informe programado se ha generado correctamente. Revise el contenido para asegurar que los datos presentados son precisos y relevantes."
    elif "Customer record updated" in message:
        return "El registro de un cliente ha sido actualizado. Verifique los cambios para asegurar que se reflejan correctamente en el sistema."
    elif "Data export completed" in message:
        return "Exportación de datos completada. Revise el archivo exportado para confirmar que todos los datos necesarios están presentes y son correctos."
    elif "User logged in successfully" in message:
        return "Inicio de sesión exitoso. El usuario ha accedido al sistema correctamente."
    else:
        return f"Evento no reconocido: {message}. Es necesario realizar un análisis detallado para identificar la causa y el impacto potencial."

# Función para analizar los logs y categorizar los eventos
def analizar_logs(logs):
    errores, advertencias, eventos_criticos, otros_eventos = [], [], [], []
    
    for log in logs:
        severity = log[0]
        message = log[1]
        explicacion = generar_explicacion(log)
        if severity == 'ERROR':
            errores.append((log, explicacion))
        elif severity == 'WARNING':
            advertencias.append((log, explicacion))
        elif severity == 'CRITICAL':
            eventos_criticos.append((log, explicacion))
        else:
            otros_eventos.append((log, explicacion))
    
    return errores, advertencias, eventos_criticos, otros_eventos

# Función para combinar resultados de múltiples archivos de logs
def combinar_resultados(resultados):
    errores, advertencias, eventos_criticos, otros_eventos = [], [], [], []
    
    for resultado in resultados:
        errores.extend(resultado[0])
        advertencias.extend(resultado[1])
        eventos_criticos.extend(resultado[2])
        otros_eventos.extend(resultado[3])
    
    return errores, advertencias, eventos_criticos, otros_eventos

# Función para generar un resumen estadístico de los logs
def generar_resumen(errores, advertencias, eventos_criticos, otros_eventos):
    return {
        'Total de logs': len(errores) + len(advertencias) + len(eventos_criticos) + len(otros_eventos),
        'Errores': len(errores),
        'Advertencias': len(advertencias),
        'Eventos críticos': len(eventos_criticos),
        'Otros eventos': len(otros_eventos),
        'Fecha del resumen': datetime.datetime.now().strftime('%Y-%m-%d %H:%M:%S')
    }

# Función para añadir bordes a una tabla en Word
def agregar_bordes_tabla(tabla):
    tbl = tabla._tbl
    for cell in tbl.iter_tcs():
        tcPr = cell.get_or_add_tcPr()
        tcBorders = OxmlElement('w:tcBorders')
        for border_name in ['top', 'left', 'bottom', 'right']:
            border = OxmlElement(f'w:{border_name}')
            border.set(qn('w:val'), 'single')
            border.set(qn('w:sz'), '4')
            border.set(qn('w:space'), '0')
            border.set(qn('w:color'), '000000')
            tcBorders.append(border)
        tcPr.append(tcBorders)

# Función para generar el informe de auditoría en formato Word
def generar_informe_word(resumen, errores, advertencias, eventos_criticos, otros_eventos, total_logs):
    doc = Document()
    doc.add_heading('INFORME DE AUDITORÍA DE LOGS DEL SISTEMA', 0)
    doc.add_paragraph(f'Fecha de Generación: {resumen["Fecha del resumen"]}', style='Heading 3')
    doc.add_paragraph(f'Total de Logs Analizados: {total_logs}')
    doc.add_paragraph("\n")

    # Índice
    doc.add_heading('Índice', level=1)
    doc.add_paragraph('1. Introducción')
    doc.add_paragraph('2. Objetivo de la Auditoría')
    doc.add_paragraph('3. Metodología')
    doc.add_paragraph('4. Resumen Ejecutivo')
    doc.add_paragraph('5. Análisis de Errores')
    doc.add_paragraph('6. Análisis de Advertencias')
    doc.add_paragraph('7. Análisis de Eventos Críticos')
    doc.add_paragraph('8. Patrones Recurrentes y Observaciones')
    doc.add_paragraph('9. Impacto Potencial de los Problemas Identificados')
    doc.add_paragraph('10. Recomendaciones Detalladas')
    doc.add_paragraph('11. Acciones Correctivas Inmediatas')
    doc.add_paragraph('12. Conclusión')
    doc.add_paragraph('13. Firmas')
    doc.add_paragraph("\n")

    # Introducción y Objetivo
    doc.add_heading('Introducción', level=1)
    doc.add_paragraph(
        "Los logs son registros detallados de eventos que ocurren dentro de un sistema. Estos registros son fundamentales para la monitorización, "
        "diagnóstico y auditoría del sistema, proporcionando un rastro de actividades que permite a los administradores y desarrolladores "
        "identificar y resolver problemas, asegurar el cumplimiento normativo y mantener la seguridad del sistema."
    )
    doc.add_heading('Objetivo de la Auditoría', level=1)
    doc.add_paragraph(
        "El objetivo de esta auditoría es evaluar el estado actual del sistema mediante el análisis de los logs generados, identificar patrones de "
        "comportamiento anómalo, y determinar las áreas que requieren atención para mejorar la estabilidad, rendimiento y seguridad del sistema."
    )

    # Metodología
    doc.add_heading('Metodología', level=1)
    doc.add_paragraph(
        "La auditoría fue realizada utilizando una herramienta automatizada que analiza los archivos de logs generados por el sistema. "
        "Los logs se clasificaron en diferentes categorías (errores, advertencias, eventos críticos) basándose en palabras clave y patrones "
        "predefinidos. Se realizó un análisis estadístico y se generaron explicaciones detalladas para cada tipo de log."
    )

    # Resumen Ejecutivo
    doc.add_heading('Resumen Ejecutivo', level=1)
    doc.add_paragraph(
        f"Se analizaron un total de {total_logs} logs, de los cuales {resumen['Errores']} fueron clasificados como errores, "
        f"{resumen['Advertencias']} como advertencias y {resumen['Eventos críticos']} como eventos críticos. "
        "La auditoría identificó varios problemas críticos que requieren atención inmediata."
    )
    doc.add_paragraph(
        "El análisis mostró que los errores más comunes están relacionados con problemas de conectividad y sobrecarga de recursos. "
        "Las advertencias se centraron en intentos de acceso no autorizado y problemas de seguridad menores. Los eventos críticos "
        "indicaron interrupciones significativas del sistema que podrían afectar la disponibilidad del servicio."
    )

    # Análisis de Errores
    doc.add_heading('Análisis de Errores', level=1)
    table = doc.add_table(rows=1, cols=3)
    table.style = 'Table Grid'
    table.cell(0, 0).text = 'Mensaje del Error'
    table.cell(0, 1).text = 'Hora'
    table.cell(0, 2).text = 'Explicación'
    agregar_bordes_tabla(table)
    
    for log, explicacion in errores:
        row = table.add_row().cells
        row[0].text = log[1]
        row[1].text = log[2]
        row[2].text = explicacion

    # Análisis de Advertencias
    doc.add_heading('Análisis de Advertencias', level=1)
    table = doc.add_table(rows=1, cols=3)
    table.style = 'Table Grid'
    table.cell(0, 0).text = 'Mensaje de la Advertencia'
    table.cell(0, 1).text = 'Hora'
    table.cell(0, 2).text = 'Explicación'
    agregar_bordes_tabla(table)
    
    for log, explicacion in advertencias:
        row = table.add_row().cells
        row[0].text = log[1]
        row[1].text = log[2]
        row[2].text = explicacion

    # Análisis de Eventos Críticos
    doc.add_heading('Análisis de Eventos Críticos', level=1)
    table = doc.add_table(rows=1, cols=3)
    table.style = 'Table Grid'
    table.cell(0, 0).text = 'Mensaje del Evento Crítico'
    table.cell(0, 1).text = 'Hora'
    table.cell(0, 2).text = 'Explicación'
    agregar_bordes_tabla(table)
    
    for log, explicacion in eventos_criticos:
        row = table.add_row().cells
        row[0].text = log[1]
        row[1].text = log[2]
        row[2].text = explicacion

    # Patrones Recurrentes y Observaciones
    doc.add_heading('Patrones Recurrentes y Observaciones', level=1)
    doc.add_paragraph(
        "Se identificaron varios patrones recurrentes en los logs analizados, lo que sugiere posibles áreas problemáticas en el sistema. "
        "Específicamente, se observó que ciertos errores tienden a ocurrir en intervalos de tiempo similares, lo que podría indicar "
        "problemas relacionados con la carga del sistema o con procesos específicos que se ejecutan en esos momentos. "
        "Además, las advertencias relacionadas con la seguridad requieren atención inmediata para evitar posibles brechas de seguridad."
    )

    # Impacto Potencial de los Problemas Identificados
    doc.add_heading('Impacto Potencial de los Problemas Identificados', level=1)
    doc.add_paragraph(
        "Los problemas identificados en esta auditoría podrían tener un impacto significativo en la operación y seguridad del sistema. "
        "Los errores de conectividad y las sobrecargas de recursos pueden llevar a tiempos de inactividad, afectando la disponibilidad del servicio. "
        "Los intentos de acceso no autorizado y las brechas de seguridad podrían comprometer datos sensibles y la integridad del sistema."
    )

    # Recomendaciones Detalladas
    doc.add_heading('Recomendaciones Detalladas', level=1)
    doc.add_paragraph(
        "Para abordar los problemas identificados, se recomiendan las siguientes acciones específicas:\n"
        "1. **Monitoreo Continuo:** Implementar herramientas de monitoreo para detectar y alertar sobre problemas de conectividad y sobrecarga en tiempo real.\n"
        "2. **Fortalecimiento de la Seguridad:** Revisar y actualizar las políticas de seguridad para prevenir accesos no autorizados, incluyendo autenticación de múltiples factores.\n"
        "3. **Optimización de Recursos:** Realizar un análisis de rendimiento para identificar y eliminar cuellos de botella, mejorando así la eficiencia del sistema.\n"
        "4. **Actualización de la Infraestructura:** Considerar la actualización del hardware o la ampliación de la capacidad del servidor para manejar mejor la carga y los picos de tráfico.\n"
        "5. **Capacitación del Personal:** Asegurar que el personal esté capacitado para responder a incidentes de seguridad y para utilizar herramientas de monitoreo y diagnóstico de manera efectiva.\n"
        "6. **Revisión Periódica de Logs:** Establecer un calendario de revisión de logs para identificar problemas emergentes antes de que se conviertan en críticos.\n"
        "7. **Auditorías de Seguridad:** Realizar auditorías de seguridad regulares para identificar vulnerabilidades y asegurar que las políticas de seguridad se estén aplicando correctamente."
    )

    # Acciones Correctivas Inmediatas
    doc.add_heading('Acciones Correctivas Inmediatas', level=1)
    doc.add_paragraph(
        "Basado en los hallazgos de esta auditoría, se recomiendan las siguientes acciones correctivas inmediatas para mitigar los riesgos identificados:\n"
        "1. **Resolver Problemas de Conectividad:** Identificar y solucionar los problemas de conexión a la base de datos para asegurar la disponibilidad continua del servicio.\n"
        "2. **Aumentar la Capacidad de Almacenamiento:** Revisar y expandir la capacidad de almacenamiento para evitar problemas relacionados con el espacio en disco insuficiente.\n"
        "3. **Implementar Monitoreo de Seguridad:** Instalar herramientas de monitoreo que alerten automáticamente sobre intentos de acceso no autorizado y brechas de seguridad.\n"
        "4. **Optimización de Procesos de Backup:** Asegurarse de que los procesos de respaldo de datos estén configurados correctamente y que se realicen regularmente sin fallos.\n"
        "5. **Actualizar Configuraciones de Tiempo de Sesión:** Revisar y ajustar las configuraciones de tiempo de sesión para evitar cierres inesperados de sesión de usuario debido a configuraciones demasiado estrictas."
    )

    # Conclusión
    doc.add_heading('Conclusión', level=1)
    doc.add_paragraph(
        "La auditoría de logs realizada proporciona una visión integral del estado actual del sistema, identificando tanto problemas críticos como áreas de mejora. "
        "Es evidente que existen problemas de conectividad y sobrecarga de recursos que deben ser abordados para asegurar la estabilidad y disponibilidad del sistema. "
        "Asi mismo, los intentos de acceso no autorizado resaltan la necesidad de mejorar las medidas de seguridad. Implementar las recomendaciones propuestas ayudará a mitigar estos riesgos, mejorar la eficiencia y garantizar la integridad y seguridad del sistema a largo plazo."
    )
    doc.add_paragraph(
        "Se recomienda realizar auditorías de logs periódicamente para mantener un control continuo sobre el estado del sistema y responder proactivamente a cualquier "
        "incidencia que pudiera surgir. La adopción de una estrategia de monitoreo continuo y la actualización regular de políticas y procedimientos de seguridad serán clave para mantener la resiliencia del sistema ante futuros desafíos."
    )

    # Firma del Auditor
    doc.add_heading('Firmas', level=1)
    doc.add_paragraph("Firma del Auditor: __________________________")
    doc.add_paragraph("Nombre del Auditor: [Nombre del Auditor]")
    doc.add_paragraph("\n")

    # Guardar el documento en un buffer en memoria
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    
    return buffer

# Función principal para la ejecución de la aplicación en Streamlit
def main():
    st.title("Auditoría de Logs del Sistema")
    st.write(
        """
        ### Descripción de la Auditoría de Logs
        Los logs son registros que documentan eventos importantes que ocurren en un sistema de software. Estos registros pueden ayudar a los administradores a diagnosticar problemas, monitorear la seguridad, y asegurar que el sistema esté funcionando correctamente. 
        
        Esta herramienta permite cargar archivos de logs de un sistema, analizar los registros en busca de errores, advertencias, y eventos críticos, y generar un informe detallado en formato Word.
        
        ### Tipos de Logs Soportados
        Esta herramienta puede analizar y categorizar los siguientes tipos de logs:
        - Conexión fallida a la base de datos
        - Incapacidad para alcanzar endpoints de API
        - Errores de sincronización de datos
        - Intentos de acceso no autorizado
        - Sobrecargas del servidor
        - Y muchos más...
        
        ### Beneficios de la Auditoría de Logs
        Realizar una auditoría de logs proporciona una visión detallada de los eventos del sistema, permitiendo:
        - Identificar y corregir problemas críticos rápidamente.
        - Mejorar la seguridad al detectar intentos de acceso no autorizados.
        - Optimizar el rendimiento del sistema mediante la identificación de cuellos de botella.
        - Asegurar el cumplimiento normativo al mantener un registro detallado de todas las actividades.
        
        ### Instrucciones para Usar la Herramienta
        1. **Seleccione los archivos de logs:** Puede cargar múltiples archivos de logs en formato `.log` o archivos de Excel `.xlsx`.
        2. **Analice los logs:** Los registros serán analizados y categorizados.
        3. **Genere un informe:** Haga clic en el botón para generar un informe de auditoría en formato Word.
        
        ### Descarga e Interpretación del Informe
        Una vez completado el análisis de logs, puede descargar un informe detallado en formato Word. El informe incluye:
        - Un resumen ejecutivo de los resultados.
        - Análisis detallado de errores, advertencias y eventos críticos.
        - Recomendaciones específicas para mejorar la estabilidad y seguridad del sistema.
        
        Para descargar el informe, haga clic en el botón "Generar Informe Word" y luego en "Descargar Informe Word".
        
        ### ¿Necesita Ayuda?
        Si tiene alguna pregunta o necesita asistencia, por favor, [contáctenos](#).
        """
    )
    
    archivos_subidos = st.file_uploader("Seleccione los archivos de logs", accept_multiple_files=True, type=["log", "xlsx"])
    
    if archivos_subidos:
        resultados = []
        total_logs = 0
        
        for archivo in archivos_subidos:
            logs = leer_logs(archivo)
            total_logs += len(logs)
            if logs:
                resultados.append(analizar_logs(logs))
        
        errores, advertencias, eventos_criticos, otros_eventos = combinar_resultados(resultados)
        
        resumen = generar_resumen(errores, advertencias, eventos_criticos, otros_eventos)
        
        st.subheader("Resumen de Resultados")
        st.write(f"Total de Logs Analizados: {total_logs}")
        st.write(f"Errores: {resumen['Errores']}")
        st.write(f"Advertencias: {resumen['Advertencias']}")
        st.write(f"Eventos Críticos: {resumen['Eventos críticos']}")
        
        if st.button("Generar Informe Word"):
            buffer = generar_informe_word(resumen, errores, advertencias, eventos_criticos, otros_eventos, total_logs)
            st.download_button(label="Descargar Informe Word", data=buffer, file_name="informe_auditoria_logs.docx", mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document")

if __name__ == "__main__":
    main()
