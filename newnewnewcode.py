import os
import datetime
from collections import Counter
from tkinter import Tk, filedialog
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH


def leer_logs(ruta_archivo):
    try:
        with open(ruta_archivo, 'r', encoding='latin-1') as file:
            return file.readlines()
    except FileNotFoundError:
        print(f"Error: El archivo {ruta_archivo} no se encontró.")
    except PermissionError:
        print(f"Error: No se tienen permisos para leer el archivo {ruta_archivo}.")
    except Exception as e:
        print(f"Error inesperado al leer el archivo {ruta_archivo}: {e}")
    return []


def generar_explicacion(log):
    # Simplificamos la lógica para buscar y explicar los errores.
    errores_explicaciones = {
        'Database connection failed': "El sistema no pudo establecer una conexión con la base de datos. "
                                      "Esto puede deberse a problemas de red, credenciales incorrectas o un fallo en el servicio de la base de datos.",
        'Unable to reach API endpoint': "El sistema no pudo comunicarse con el punto final de API. "
                                        "Esto puede deberse a un servidor API caído o problemas de red.",
        'Failed to back up database': "El respaldo de la base de datos no se pudo completar. "
                                      "Esto puede deberse a falta de espacio en disco o problemas de permisos.",
        'High memory usage detected': "El sistema está utilizando una cantidad inusualmente alta de memoria, lo que podría llevar a un rendimiento lento o incluso a un bloqueo del sistema.",
        'Disk space low': "El espacio en disco está llegando al límite, lo que podría impedir la capacidad del sistema para realizar operaciones de escritura.",
        'Slow response time': "El tiempo de respuesta del sistema es más lento de lo esperado, lo que podría ser causado por una carga excesiva o cuellos de botella en el procesamiento.",
        'System outage detected': "Se ha detectado una interrupción del sistema, lo que podría deberse a fallas en el hardware o problemas en la red.",
        'Security breach detected': "Se ha detectado una posible brecha de seguridad, lo que podría indicar accesos no autorizados o intentos de intrusión.",
        'Application crash': "Una aplicación del sistema se ha bloqueado inesperadamente, posiblemente debido a errores en el código o conflictos de recursos."
    }

    for clave, explicacion in errores_explicaciones.items():
        if clave in log:
            return explicacion

    return "Este evento registrado requiere una revisión detallada para comprender completamente su impacto."


def analizar_logs(logs):
    errores, advertencias, eventos_criticos, otros_eventos = [], [], [], []

    for log in logs:
        explicacion = generar_explicacion(log)
        if 'ERROR' in log:
            errores.append((log, explicacion))
        elif 'WARNING' in log:
            advertencias.append((log, explicacion))
        elif 'CRITICAL' in log:
            eventos_criticos.append((log, explicacion))
        else:
            otros_eventos.append((log, explicacion))
    
    return errores, advertencias, eventos_criticos, otros_eventos


def combinar_resultados(resultados):
    errores, advertencias, eventos_criticos, otros_eventos = [], [], [], []

    for resultado in resultados:
        errores.extend(resultado[0])
        advertencias.extend(resultado[1])
        eventos_criticos.extend(resultado[2])
        otros_eventos.extend(resultado[3])

    return errores, advertencias, eventos_criticos, otros_eventos


def obtener_descripcion(log):
    try:
        return ' '.join(log.split(' ')[3:])
    except IndexError:
        return 'Descripción no disponible'


def obtener_hora(log):
    try:
        return log.split(' ')[1]
    except IndexError:
        return 'Desconocido'


def generar_objetivo_prueba(errores, advertencias, eventos_criticos):
    if eventos_criticos:
        return f"Investigar {len(eventos_criticos)} eventos críticos que podrían comprometer la estabilidad del sistema."
    elif errores:
        return f"Analizar los {len(errores)} errores registrados para mejorar la fiabilidad del sistema."
    elif advertencias:
        return f"Revisar las {len(advertencias)} advertencias para prevenir futuros errores."
    else:
        return "Asegurar que no existen problemas críticos que afecten el rendimiento del sistema."


def generar_resumen(errores, advertencias, eventos_criticos, otros_eventos):
    return {
        'Total de logs': len(errores) + len(advertencias) + len(eventos_criticos) + len(otros_eventos),
        'Errores': len(errores),
        'Advertencias': len(advertencias),
        'Eventos críticos': len(eventos_criticos),
        'Otros eventos': len(otros_eventos),
        'Errores más comunes': Counter([obtener_descripcion(log[0]) for log in errores]).most_common(5),
        'Advertencias más comunes': Counter([obtener_descripcion(log[0]) for log in advertencias]).most_common(5),
        'Eventos críticos más comunes': Counter([obtener_descripcion(log[0]) for log in eventos_criticos]).most_common(5),
        'Frecuencia de errores por hora': Counter([obtener_hora(log[0]) for log in errores]),
        'Frecuencia de advertencias por hora': Counter([obtener_hora(log[0]) for log in advertencias]),
        'Frecuencia de eventos críticos por hora': Counter([obtener_hora(log[0]) for log in eventos_criticos]),
        'Fecha del resumen': datetime.datetime.now().strftime('%Y-%m-%d %H:%M:%S')
    }


def agregar_encabezado_tabla(tabla, encabezados):
    fila = tabla.rows[0]
    for idx, encabezado in enumerate(encabezados):
        celda = fila.cells[idx]
        celda.text = encabezado
        parrafo = celda.paragraphs[0]
        parrafo.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = parrafo.runs[0]
        run.bold = True


def agregar_encabezado(doc, texto):
    parrafo = doc.add_paragraph(texto)
    parrafo.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = parrafo.runs[0]
    run.bold = True
    run.font.size = Pt(14)


def ajustar_ancho_columnas(tabla):
    for row in tabla.rows:
        for cell in row.cells:
            cell.width = Inches(2)
            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER


def generar_informe_word(resumen, archivo_salida, objetivo_prueba, errores, advertencias, eventos_criticos, total_logs):
    try:
        doc = Document()

        # Portada del documento
        agregar_encabezado(doc, "INFORME DE AUDITORÍA DE LOGS DEL SISTEMA")
        agregar_encabezado(doc, "Fecha: " + resumen['Fecha del resumen'])
        agregar_encabezado(doc, "Auditor: [Nombre del Auditor]")
        agregar_encabezado(doc, "Empresa: [Nombre de la Empresa]")
        doc.add_paragraph("\n\n")

        # Índice del documento
        doc.add_heading("Índice", level=1)
        doc.add_paragraph("1. Introducción")
        doc.add_paragraph("2. Objetivo de la Auditoría")
        doc.add_paragraph("3. Resumen de Resultados")
        doc.add_paragraph("4. Análisis de Errores Más Comunes")
        doc.add_paragraph("5. Análisis de Advertencias Más Comunes")
        doc.add_paragraph("6. Análisis de Eventos Críticos Más Comunes")
        doc.add_paragraph("7. Distribución Temporal de Eventos")
        doc.add_paragraph("8. Detalle de Errores y Problemas Encontrados")
        doc.add_paragraph("9. Conclusión")
        doc.add_page_break()

        # Introducción
        doc.add_heading("1. Introducción", level=1)
        doc.add_paragraph(f"Este informe de auditoría de logs se basa en el análisis de {total_logs} registros de eventos del sistema. "
                          "Los logs del sistema son una fuente crucial de información que permite identificar y diagnosticar posibles problemas "
                          "y asegurar que el sistema esté funcionando de manera óptima.")

        # Objetivo de la auditoría
        doc.add_heading("2. Objetivo de la Auditoría", level=1)
        doc.add_paragraph(objetivo_prueba)

        # Resumen de resultados
        doc.add_heading("3. Resumen de Resultados", level=1)
        doc.add_paragraph(f"Total de Logs Analizados: {total_logs}")
        doc.add_paragraph(f"Total de Errores: {resumen['Errores']}")
        doc.add_paragraph(f"Total de Advertencias: {resumen['Advertencias']}")
        doc.add_paragraph(f"Total de Eventos Críticos: {resumen['Eventos críticos']}")
        doc.add_paragraph(f"Total de Otros Eventos: {resumen['Otros eventos']}")

        # Análisis de errores más comunes
        doc.add_heading("4. Análisis de Errores Más Comunes", level=1)
        if resumen.get('Errores más comunes'):
            doc.add_paragraph("A continuación se detallan los errores más comunes encontrados:")
            tabla = doc.add_table(rows=1, cols=2)
            tabla.style = 'Table Grid'
            agregar_encabezado_tabla(tabla, ["Descripción del Error", "Ocurrencias"])
            for error in resumen.get('Errores más comunes', []):
                fila = tabla.add_row().cells
                fila[0].text = error[0]
                fila[1].text = str(error[1])
            ajustar_ancho_columnas(tabla)
        else:
            doc.add_paragraph("No se encontraron errores comunes.")

        # Análisis de advertencias más comunes
        doc.add_heading("5. Análisis de Advertencias Más Comunes", level=1)
        if resumen.get('Advertencias más comunes'):
            doc.add_paragraph("A continuación se detallan las advertencias más comunes encontradas:")
            tabla = doc.add_table(rows=1, cols=2)
            tabla.style = 'Table Grid'
            agregar_encabezado_tabla(tabla, ["Descripción de la Advertencia", "Ocurrencias"])
            for advertencia in resumen.get('Advertencias más comunes', []):
                fila = tabla.add_row().cells
                fila[0].text = advertencia[0]
                fila[1].text = str(advertencia[1])
            ajustar_ancho_columnas(tabla)
        else:
            doc.add_paragraph("No se encontraron advertencias comunes.")

        # Análisis de eventos críticos más comunes
        doc.add_heading("6. Análisis de Eventos Críticos Más Comunes", level=1)
        if resumen.get('Eventos críticos más comunes'):
            doc.add_paragraph("A continuación se detallan los eventos críticos más comunes encontrados:")
            tabla = doc.add_table(rows=1, cols=2)
            tabla.style = 'Table Grid'
            agregar_encabezado_tabla(tabla, ["Descripción del Evento Crítico", "Ocurrencias"])
            for evento in resumen.get('Eventos críticos más comunes', []):
                fila = tabla.add_row().cells
                fila[0].text = evento[0]
                fila[1].text = str(evento[1])
            ajustar_ancho_columnas(tabla)
        else:
            doc.add_paragraph("No se encontraron eventos críticos comunes.")

        # Distribución temporal de eventos
        doc.add_heading("7. Distribución Temporal de Eventos", level=1)
        doc.add_paragraph("En esta sección se analiza la distribución de errores, advertencias y eventos críticos a lo largo del tiempo, lo que puede ayudar a identificar patrones recurrentes o periodos de mayor actividad o problemas en el sistema.")
        
        doc.add_paragraph("Frecuencia de Errores por Hora:")
        if resumen.get('Frecuencia de errores por hora'):
            tabla = doc.add_table(rows=1, cols=2)
            tabla.style = 'Table Grid'
            agregar_encabezado_tabla(tabla, ["Hora", "Errores"])
            for hora, frecuencia in sorted(resumen.get('Frecuencia de errores por hora', {}).items()):
                fila = tabla.add_row().cells
                fila[0].text = hora
                fila[1].text = str(frecuencia)
            ajustar_ancho_columnas(tabla)
        else:
            doc.add_paragraph("No se registraron errores en el periodo analizado.")

        doc.add_paragraph("Frecuencia de Advertencias por Hora:")
        if resumen.get('Frecuencia de advertencias por hora'):
            tabla = doc.add_table(rows=1, cols=2)
            tabla.style = 'Table Grid'
            agregar_encabezado_tabla(tabla, ["Hora", "Advertencias"])
            for hora, frecuencia in sorted(resumen.get('Frecuencia de advertencias por hora', {}).items()):
                fila = tabla.add_row().cells
                fila[0].text = hora
                fila[1].text = str(frecuencia)
            ajustar_ancho_columnas(tabla)
        else:
            doc.add_paragraph("No se registraron advertencias en el periodo analizado.")

        doc.add_paragraph("Frecuencia de Eventos Críticos por Hora:")
        if resumen.get('Frecuencia de eventos críticos por hora'):
            tabla = doc.add_table(rows=1, cols=2)
            tabla.style = 'Table Grid'
            agregar_encabezado_tabla(tabla, ["Hora", "Eventos Críticos"])
            for hora, frecuencia in sorted(resumen.get('Frecuencia de eventos críticos por hora', {}).items()):
                fila = tabla.add_row().cells
                fila[0].text = hora
                fila[1].text = str(frecuencia)
            ajustar_ancho_columnas(tabla)
        else:
            doc.add_paragraph("No se registraron eventos críticos en el periodo analizado.")

        # Detalle de errores y problemas encontrados
        doc.add_heading("8. Detalle de Errores y Problemas Encontrados", level=1)
        doc.add_paragraph("A continuación se presentan los detalles de los errores, advertencias y eventos críticos encontrados, junto con una explicación detallada del posible impacto y recomendaciones:")

        if errores:
            doc.add_heading("Errores:", level=2)
            for log, explicacion in errores:
                doc.add_paragraph(f"Log: {log}")
                doc.add_paragraph(f"Explicación: {explicacion}")
        else:
            doc.add_paragraph("No se encontraron errores en los logs analizados.")

        if advertencias:
            doc.add_heading("Advertencias:", level=2)
            for log, explicacion in advertencias:
                doc.add_paragraph(f"Log: {log}")
                doc.add_paragraph(f"Explicación: {explicacion}")
        else:
            doc.add_paragraph("No se encontraron advertencias en los logs analizados.")

        if eventos_criticos:
            doc.add_heading("Eventos Críticos:", level=2)
            for log, explicacion in eventos_criticos:
                doc.add_paragraph(f"Log: {log}")
                doc.add_paragraph(f"Explicación: {explicacion}")
        else:
            doc.add_paragraph("No se encontraron eventos críticos en los logs analizados.")

        # Conclusión
        doc.add_heading("9. Conclusión", level=1)
        doc.add_paragraph("""
        A partir del análisis realizado en los logs del sistema, se han identificado las siguientes áreas de mejora y recomendaciones:
        """)
        if eventos_criticos:
            doc.add_paragraph(f"1. Se han detectado {len(eventos_criticos)} eventos críticos que requieren atención inmediata para garantizar la estabilidad del sistema.")
        if errores:
            doc.add_paragraph(f"2. Se registraron {len(errores)} errores en el sistema. Es importante abordar estos errores para mejorar la fiabilidad.")
        if advertencias:
            doc.add_paragraph(f"3. Se encontraron {len(advertencias)} advertencias. Es recomendable revisar estas advertencias para prevenir posibles problemas futuros.")
        if not (errores or advertencias or eventos_criticos):
            doc.add_paragraph("No se encontraron problemas significativos en los logs analizados.")
        
        doc.add_paragraph("""
        Se recomienda implementar un sistema de monitoreo continuo para detectar y corregir problemas en tiempo real, así como realizar auditorías periódicas para asegurar el correcto funcionamiento del sistema.
        """)

        # Guardar el informe en un archivo .docx
        doc.save(archivo_salida)
    except PermissionError:
        print(f"Error: No se pudo escribir en el archivo {archivo_salida} debido a permisos insuficientes.")
    except Exception as e:
        print(f"Error inesperado al generar el informe: {e}")


def seleccionar_archivos():
    Tk().withdraw()
    archivos = filedialog.askopenfilenames(
        title="Seleccione los archivos de logs",
        filetypes=[("Archivos de log", "*.log"), ("Todos los archivos", "*.*")]
    )
    return list(archivos)


def obtener_carpeta_documentos():
    return os.path.expanduser("~/Documents")


def main():
    print("Auditoría de Logs del Sistema")

    rutas_archivos = seleccionar_archivos()

    if rutas_archivos:
        resultados = []
        total_logs = 0

        for ruta in rutas_archivos:
            logs = leer_logs(ruta)
            total_logs += len(logs)
            if logs:
                resultados.append(analizar_logs(logs))
        
        print(f"Total de logs a analizar: {total_logs}")

        errores, advertencias, eventos_criticos, otros_eventos = combinar_resultados(resultados)

        resumen = generar_resumen(errores, advertencias, eventos_criticos, otros_eventos)

        objetivo_prueba = generar_objetivo_prueba(errores, advertencias, eventos_criticos)

        carpeta_documentos = obtener_carpeta_documentos()
        archivo_salida = os.path.join(carpeta_documentos, 'informe_auditoria_logs.docx')
        generar_informe_word(resumen, archivo_salida, objetivo_prueba, errores, advertencias, eventos_criticos, total_logs)
        print(f"\nInforme de auditoría generado: {archivo_salida}")

        os.startfile(archivo_salida)
    else:
        print("No se seleccionaron archivos de logs.")


if __name__ == "__main__":
    main()
